"""
core/template_service.py
========================
Servicio de plantillas y configuraciones por formato.

Responsabilidades:
  1. Cargar la configuración JSON de un formato específico
  2. Resolver los placeholders contra la BD y la sesión activa:
       {{RESPONSABLES}} → nombre_completo del usuario en sesión
       {{PUESTO}}       → nombre_puesto del usuario en sesión
       {{FIRMA_DIGITAL}} → hash corto de la firma del usuario
  3. Aplicar los valores resueltos al documento antes de firmarlo

Uso:
    svc = TemplateService(db, user)
    valores = svc.resolver("F09_P_P_SGI_03", firma_hash="a4b7…cda1")
    # → {"{{RESPONSABLES}}": "Roberto Enríquez", "{{PUESTO}}": "Auxiliar ISO", ...}
"""
import json
from pathlib import Path
from core.hash_service import HashService

CONFIGS_DIR = Path(__file__).parent.parent / "templates" / "configuraciones"
FORMATOS_DIR = Path(__file__).parent.parent / "templates" / "formatos"


class TemplateService:

    def __init__(self, db, user):
        self._db   = db
        self._user = user
        self._h    = HashService()

    # ── Cargar configuración ──────────────────────────────────────────────────

    def cargar_config(self, formato_id: str) -> dict:
        """
        Carga el JSON de configuración para un formato dado.
        Lanza FileNotFoundError si no existe.
        """
        path = CONFIGS_DIR / f"configuracion_{formato_id}.json"
        if not path.exists():
            raise FileNotFoundError(
                f"No existe configuración para el formato '{formato_id}'.\n"
                f"Ruta buscada: {path}"
            )
        with open(path, encoding="utf-8") as f:
            return json.load(f)

    def listar_formatos(self) -> list[dict]:
        """Devuelve lista de formatos configurados."""
        resultado = []
        for cfg_file in CONFIGS_DIR.glob("configuracion_*.json"):
            try:
                with open(cfg_file, encoding="utf-8") as f:
                    data = json.load(f)
                resultado.append({
                    "formato_id"     : data.get("formato_id", cfg_file.stem),
                    "nombre_formato" : data.get("nombre_formato", ""),
                    "extension"      : data.get("extension_esperada", ""),
                    "config_path"    : str(cfg_file),
                })
            except Exception:
                pass
        return resultado

    def plantilla_path(self, formato_id: str) -> Path | None:
        """Devuelve la ruta del archivo de plantilla si existe."""
        config  = self.cargar_config(formato_id)
        ext     = config.get("extension_esperada", "")
        pattern = f"{formato_id}{ext}"
        path    = FORMATOS_DIR / pattern
        return path if path.exists() else None

    # ── Resolver placeholders ─────────────────────────────────────────────────

    def resolver(
        self,
        formato_id: str,
        firma_hash: str | None = None,
    ) -> dict[str, str]:
        """
        Devuelve un dict {placeholder: valor_resuelto} para el formato dado.

        Los valores se obtienen de la BD comparados con la sesión activa:
          {{RESPONSABLES}} → user.nombre_completo
          {{PUESTO}}       → user.nombre_puesto
          {{FIRMA_DIGITAL}}→ hash corto (4+4 chars) de la firma más reciente
                             o del firma_hash pasado explícitamente
        """
        config       = self.cargar_config(formato_id)
        placeholders = config.get("placeholders", {})
        resultado    = {}

        for ph, cfg in placeholders.items():
            modo = cfg.get("modo_insercion", "inline")

            if ph == "{{RESPONSABLES}}":
                resultado[ph] = self._user.nombre_completo

            elif ph == "{{PUESTO}}":
                resultado[ph] = self._user.nombre_puesto

            elif ph == "{{FIRMA_DIGITAL}}":
                # Usar el hash pasado o buscar el más reciente en BD
                if firma_hash:
                    short = self._h.short_hash(firma_hash)
                else:
                    short = self._ultimo_hash_corto()
                resultado[ph] = f"#{short}"

        return resultado

    def _ultimo_hash_corto(self) -> str:
        """Obtiene el hash corto de la firma más reciente del usuario."""
        sigs = self._db.sig_repo.get_by_user(self._user.id_user)
        if not sigs:
            return "——————"
        last = sigs[-1]
        full = last.get("firma_hash", "") if isinstance(last, dict) \
               else getattr(last, "firma_hash", "")
        return self._h.short_hash(full)

    # ── Aplicar al documento ──────────────────────────────────────────────────

    def aplicar_a_xlsx(
        self,
        filepath: str,
        formato_id: str,
        firma_hash: str | None = None,
    ) -> int:
        """
        Reemplaza placeholders en un archivo Excel (.xlsx).
        Expande filas si hay múltiples responsables.
        Devuelve el número de reemplazos realizados.
        """
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        config    = self.cargar_config(formato_id)
        estilos   = config.get("estilos", {})
        valores   = self.resolver(formato_id, firma_hash=firma_hash)

        wb      = openpyxl.load_workbook(filepath)
        cambios = 0

        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    val = str(cell.value or "").strip()
                    if val in valores:
                        nuevo = valores[val]
                        cell.value = nuevo
                        _aplicar_estilo(cell, val, estilos)
                        cambios += 1

        wb.save(filepath)
        return cambios

    def aplicar_a_docx(
        self,
        filepath: str,
        formato_id: str,
        firma_hash: str | None = None,
    ) -> int:
        """
        Reemplaza placeholders en un archivo Word (.docx).
        Devuelve el número de reemplazos realizados.
        """
        from docx import Document
        config  = self.cargar_config(formato_id)
        valores = self.resolver(formato_id, firma_hash=firma_hash)
        doc     = Document(filepath)
        cambios = 0

        def reemplazar_en_texto(texto: str) -> tuple[str, bool]:
            original = texto
            for ph, val in valores.items():
                if ph in texto:
                    texto = texto.replace(ph, val)
            return texto, texto != original

        for para in doc.paragraphs:
            for run in para.runs:
                nuevo, cambio = reemplazar_en_texto(run.text)
                if cambio:
                    run.text = nuevo
                    cambios += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            nuevo, cambio = reemplazar_en_texto(run.text)
                            if cambio:
                                run.text = nuevo
                                cambios += 1

        doc.save(filepath)
        return cambios


# ── helpers internos ──────────────────────────────────────────────────────────

def _aplicar_estilo(cell, placeholder: str, estilos: dict):
    from openpyxl.styles import Font, PatternFill, Alignment
    is_hash = placeholder == "{{FIRMA_DIGITAL}}"
    color_hex = estilos.get("color_hash" if is_hash else "color_datos", "1A1A2E")
    size      = estilos.get("font_size_hash" if is_hash else "font_size_datos", 9)
    fname     = estilos.get("font_name", "Calibri")

    cell.font      = Font(name=fname, size=size,
                          color=color_hex.lstrip("#"),
                          bold=(not is_hash))
    cell.alignment = Alignment(vertical="center")
