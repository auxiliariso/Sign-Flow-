"""
signers/docx_signer.py  (v2.1)
==============================
FIXES:
  - Detección de placeholders corregida (antes no matcheaba {{SIGNFLOW}})
  - La firma visible solo muestra: Nombre + Hash corto (4+4)
  - No sobrescribe firmas previas; cada firma agrega un nuevo bloque
  - Busca en párrafos Y en celdas de tablas (documentos ISO)
  - Conserva merges, bordes y estilos originales
"""
import io
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.signer_base import SignerBase, SignaturePayload, SignResult
from core.hash_service import HashService

_NAVY  = RGBColor(0x1F, 0x38, 0x64)
_MID   = RGBColor(0x2E, 0x50, 0x90)
_LIGHT = RGBColor(0xD6, 0xE4, 0xF0)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_DARK  = RGBColor(0x1A, 0x1A, 0x2E)
_GREY  = RGBColor(0x6B, 0x7A, 0x99)

# ── Keywords a detectar ── (en minúsculas para comparación) ──────────────────
PLACEHOLDER_EXACT = {
    "{{signflow}}", "{{firma_responsable}}", "{{multifirma}}", "{{responsables}}",
}
LABEL_KEYWORDS = {
    "firma:", "responsable:", "responsables:", "nombre:", "nombre y firma:",
    "revisó:", "autorizo:", "autorizó:", "aprobó:", "elaboró:", "verificó:",
}


def _matches(text: str) -> bool:
    """True si el texto contiene algún placeholder o keyword de firma."""
    t = text.strip().lower()
    # Placeholder exacto (e.g. "{{SIGNFLOW}}")
    for ph in PLACEHOLDER_EXACT:
        if ph in t:
            return True
    # Label al inicio de celda/párrafo (e.g. "Firma:  ___")
    for kw in LABEL_KEYWORDS:
        if t.startswith(kw) or t == kw.rstrip(":"):
            return True
    return False


class DocxSigner(SignerBase):

    def sign(
        self,
        doc_path: str,
        output_path: str,
        payload: SignaturePayload,
        all_previous: list[SignaturePayload],
    ) -> SignResult:
        doc = Document(doc_path)
        inserted = self._try_insert_in_table(doc, payload)
        if not inserted:
            inserted = self._try_insert_in_paragraph(doc, payload)
        if not inserted:
            self._append_signature_section(doc, payload, is_first=len(all_previous) == 0)

        doc.save(output_path)
        return SignResult(
            output_path=output_path,
            firma_hash=payload.firma_hash,
            documento_hash_post=HashService().file_hash(output_path),
        )

    def detect_signature_zones(self, doc_path: str) -> list[dict]:
        doc = Document(doc_path)
        zones = []
        for i, p in enumerate(doc.paragraphs):
            if _matches(p.text):
                zones.append({"type": "paragraph", "index": i, "text": p.text.strip()})
        for t_idx, table in enumerate(doc.tables):
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    if _matches(cell.text):
                        zones.append({"type": "cell", "table": t_idx,
                                      "row": r, "col": c, "text": cell.text.strip()})
        return zones

    # ── Inserción en tabla (documentos ISO) ───────────────────────────────────

    def _try_insert_in_table(self, doc: Document, payload: SignaturePayload) -> bool:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if _matches(cell.text):
                        self._fill_cell(cell, payload)
                        return True
        return False

    def _fill_cell(self, cell, payload: SignaturePayload):
        """
        Reemplaza el contenido de la celda con la firma.
        Solo muestra: nombre + hash corto.
        """
        # Limpiar celda
        for p in cell.paragraphs:
            for run in p.runs:
                run.text = ""

        p0 = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

        # Nombre (negrita)
        r1 = p0.add_run(payload.nombre_completo)
        r1.bold = True
        r1.font.color.rgb = _DARK
        r1.font.size = Pt(9)

        # Hash corto en línea siguiente
        p1 = cell.add_paragraph()
        r2 = p1.add_run(f"#{payload.firma_hash_short}")
        r2.bold = False
        r2.font.color.rgb = _GREY
        r2.font.size = Pt(7)

    # ── Inserción en párrafo ──────────────────────────────────────────────────

    def _try_insert_in_paragraph(self, doc: Document, payload: SignaturePayload) -> bool:
        for para in doc.paragraphs:
            if _matches(para.text):
                # Preservar el label original ("Firma:") y agregar datos a la derecha
                original_text = para.text.strip()
                for run in para.runs:
                    run.text = ""
                # Reconstruir: label + datos
                label = original_text if original_text.endswith(":") else original_text + ":"
                r_label = para.add_run(f"{label}  ")
                r_label.bold = True
                r_label.font.color.rgb = _NAVY
                r_label.font.size = Pt(9)

                r_name = para.add_run(payload.nombre_completo)
                r_name.bold = False
                r_name.font.color.rgb = _DARK
                r_name.font.size = Pt(9)

                r_hash = para.add_run(f"  #{payload.firma_hash_short}")
                r_hash.font.color.rgb = _GREY
                r_hash.font.size = Pt(7)
                return True
        return False

    # ── Bloque al final (sin placeholder) ────────────────────────────────────

    def _append_signature_section(
        self, doc: Document, payload: SignaturePayload, is_first: bool
    ):
        if is_first:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(20)
            r = sep.add_run("─" * 70)
            r.font.color.rgb = _NAVY
            r.font.size = Pt(8)

            title = doc.add_paragraph()
            rt = title.add_run("✦  FIRMAS DIGITALES — SignFlow")
            rt.bold = True
            rt.font.color.rgb = _NAVY
            rt.font.size = Pt(11)

        # Tarjeta: tabla 1×2 (datos | QR)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"

        left  = tbl.cell(0, 0)
        right = tbl.cell(0, 1)
        left.width  = Emu(5_600_000)
        right.width = Emu(1_100_000)

        lp = left.paragraphs[0]

        def add_line(txt, bold=False, color=_DARK, size=Pt(9)):
            p = left.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(txt)
            run.bold = bold
            run.font.color.rgb = color
            run.font.size = size

        # Header de la tarjeta
        rh = lp.add_run(f"  {payload.validation_id}")
        rh.bold = True
        rh.font.color.rgb = _WHITE
        rh.font.size = Pt(9)
        _set_cell_bg(left, "1F3864")

        # Solo nombre + hash corto (requisito del usuario)
        add_line(payload.nombre_completo, bold=True, color=_DARK)
        add_line(f"#{payload.firma_hash_short}", bold=False,
                 color=_GREY, size=Pt(8))
        add_line(f"{payload.fecha}  {payload.hora}", color=_MID, size=Pt(8))

        # QR
        if payload.qr_image_bytes:
            rp = right.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = rp.add_run()
            run.add_picture(io.BytesIO(payload.qr_image_bytes), width=Inches(0.75))

        doc.add_paragraph()  # espacio entre firmas


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)
