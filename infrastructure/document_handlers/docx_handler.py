"""
Handler - Firma en documentos Word (.docx)
Inserta un bloque de firma al final del documento usando python-docx.
"""
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.interfaces.repositories import IDocumentSigner

BRAND_COLOR = RGBColor(0x1F, 0x38, 0x64)   # #1F3864


class DocxSigner(IDocumentSigner):
    def sign(
        self,
        doc_path: str,
        output_path: str,
        nombre_completo: str,
        nombre_puesto: str,
        firma_hash: str,
        fecha: str,
        hora: str,
    ) -> str:
        shutil.copy2(doc_path, output_path)
        doc = Document(output_path)

        # Separador
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(18)
        run_sep = sep.add_run("─" * 60)
        run_sep.font.color.rgb = BRAND_COLOR
        run_sep.font.size = Pt(9)

        # Bloque de firma
        _signed_block(doc, nombre_completo, nombre_puesto, firma_hash, fecha, hora)

        doc.save(output_path)
        return output_path


def _signed_block(doc, nombre, puesto, hash_val, fecha, hora):
    def add_line(label: str, value: str, bold_label: bool = True):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if bold_label:
            r_label = p.add_run(f"{label}: ")
            r_label.bold = True
            r_label.font.color.rgb = BRAND_COLOR
            r_label.font.size = Pt(9)
        r_val = p.add_run(value)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = title.add_run("✦ Documento Firmado Digitalmente")
    r_title.bold = True
    r_title.font.color.rgb = BRAND_COLOR
    r_title.font.size = Pt(10)

    add_line("Firmado por", nombre)
    add_line("Cargo", puesto)
    add_line("Fecha", fecha)
    add_line("Hora (UTC)", hora)
    add_line("Hash de verificación", hash_val)
